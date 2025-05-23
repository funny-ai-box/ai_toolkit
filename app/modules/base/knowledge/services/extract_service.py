# app/modules/base/knowledge/services/extract_service.py
import logging
import io
import tempfile
import httpx # 使用 httpx 进行异步 HTTP 请求
from pathlib import Path
from typing import Protocol, runtime_checkable, Optional
import re

# 导入需要的解析库 (确保已安装: pip install python-docx python-pdfminer.six html5lib beautifulsoup4)
# 注意：.doc 文件处理已移除
try:
    import docx # python-docx for .docx
except ImportError:
    docx = None
    logging.warning("python-docx 未安装，无法解析 .docx 文件。请运行: pip install python-docx")


try:
    from pdfminer.high_level import extract_text as pdf_extract_text # <--- 方案一
    pdfminer_available = True
except ImportError:
    pdf_extract_text = None
    pdfminer_available = False
    logging.warning("pdfminer.six high_level 无法导入。请确保 pdfminer.six 已正确安装在当前环境。")

try:
    from bs4 import BeautifulSoup # 用于 HTML 解析
except ImportError:
    BeautifulSoup = None
    logging.warning("beautifulsoup4 未安装，无法优雅地解析 .html 文件。请运行: pip install beautifulsoup4 html5lib")


from app.core.config.settings import settings
from app.core.exceptions import BusinessException, NotSupportedException
from app.core.storage.base import StorageProviderType # 从 core 导入

logger = logging.getLogger(__name__)

# --- 定义协议 (接口) ---
@runtime_checkable
class IDocumentExtractService(Protocol):
    """文档内容提取服务的接口协议"""

    async def extract_file_content_async(
        self,
        document_url: str,
        original_filename: str, # 需要原始文件名来判断类型
        # storage_provider: StorageProviderType # 通常本地服务不需要知道存储类型
    ) -> str:
        """
        提取文件内容。

        Args:
            document_url: 文件的可访问 URL (本地或远程)。
            original_filename: 原始文件名 (用于判断文件类型)。

        Returns:
            提取出的文本内容。
        """
        ...

    async def extract_web_content_async(self, url: str) -> str:
        """提取网页内容"""
        ...


# --- 实现类 ---
class DocumentExtractService(IDocumentExtractService):
    """文档内容提取服务实现"""

    def __init__(self):
        # 创建异步 HTTP 客户端 (推荐使用 httpx)
        # 可以配置超时、代理等
        self._http_client = httpx.AsyncClient(
            headers={"User-Agent": "AIToolkit/1.0 (Python HttpX Client)"},
            timeout=60.0, # 设置默认超时 60 秒
            follow_redirects=True # 允许重定向
        )
        # logger.debug("DocumentExtractService 初始化完成。")

    async def _download_file(self, url: str) -> bytes:
        """使用 httpx 异步下载文件内容"""
        try:
            response = await self._http_client.get(url)
            response.raise_for_status() # 如果状态码不是 2xx，则抛出异常
            return await response.aread() # 读取响应体为 bytes
        except httpx.RequestError as e:
            logger.error(f"下载文件时请求错误: {e.request.url!r} - {e}")
            raise BusinessException(f"无法下载文件 (请求错误): {url}", code=500) from e
        except httpx.HTTPStatusError as e:
            logger.error(f"下载文件时 HTTP 状态错误: {e.request.url!r} - Status {e.response.status_code}")
            raise BusinessException(f"无法下载文件 (HTTP {e.response.status_code}): {url}", code=500) from e
        except Exception as e:
            logger.error(f"下载文件时发生未知错误: {url} - {e}")
            raise BusinessException(f"下载文件时出错: {url}", code=500) from e

    async def extract_file_content_async(
        self,
        document_url: str,
        original_filename: str,
        # storage_provider: StorageProviderType
    ) -> str:
        """提取文件内容"""
        logger.info(f"开始提取文件内容: URL='{document_url}', Filename='{original_filename}'")
        extension = Path(original_filename).suffix.lower()
        content = ""

        try:
            # 1. 获取文件内容 (bytes)
            # 对本地文件路径进行特殊处理 (如果 URL 是 file:// 或判断是本地路径)
            # 简单的判断方法可能不完全可靠，需要根据部署情况调整
            is_local = "localhost" in document_url or "127.0.0.1" in document_url or document_url.startswith("file://")
            file_bytes: Optional[bytes] = None

            if is_local and document_url.startswith(settings.LOCAL_STORAGE_BASE_URL):
                 # 如果是配置的本地存储 URL，尝试转换为本地路径
                 relative_path = document_url[len(settings.LOCAL_STORAGE_BASE_URL):].lstrip('/')
                 local_path = Path(settings.LOCAL_STORAGE_PATH) / relative_path
                 if local_path.is_file():
                     logger.debug(f"识别为本地文件，路径: {local_path}")
                     try:
                         # 使用异步方式读取本地文件 (如果需要高性能)
                         # import aiofiles
                         # async with aiofiles.open(local_path, mode='rb') as f:
                         #     file_bytes = await f.read()
                         # 同步读取对于本地文件通常足够快
                         with open(local_path, "rb") as f:
                              file_bytes = f.read()
                     except OSError as e:
                          logger.error(f"读取本地文件失败: {local_path} - {e}")
                          raise BusinessException(f"无法读取本地文件: {local_path}", code=500) from e
                 else:
                     logger.warning(f"本地文件 URL 对应的路径不存在: {local_path}, 将尝试通过 HTTP 下载。")
                     file_bytes = await self._download_file(document_url) # 回退到 HTTP 下载
            else:
                 # 认为是远程文件或无法映射的本地 URL，通过 HTTP 下载
                 logger.debug(f"将文件视为远程文件，通过 HTTP 下载: {document_url}")
                 file_bytes = await self._download_file(document_url)

            if not file_bytes:
                 raise BusinessException("未能获取文件内容", code=500)

            # 2. 根据扩展名解析内容
            file_stream = io.BytesIO(file_bytes) # 将 bytes 转为内存流

            if extension == ".txt":
                # 尝试多种编码解码
                encodings_to_try = ['utf-8', 'gbk', 'gb2312', 'latin-1']
                decoded = False
                for enc in encodings_to_try:
                    try:
                        content = file_bytes.decode(enc)
                        logger.debug(f"文本文件使用 {enc} 解码成功。")
                        decoded = True
                        break
                    except UnicodeDecodeError:
                        continue
                if not decoded:
                     logger.warning(f"无法确定文本文件的编码，将尝试忽略错误解码。")
                     content = file_bytes.decode('utf-8', errors='ignore')

            elif extension == ".html" or extension == ".htm":
                content = self._extract_html_content(file_bytes)

            elif extension == ".docx":
                if docx:
                    content = self._extract_docx_content(file_stream)
                else:
                    raise NotSupportedException("未安装 python-docx，无法解析 .docx 文件。")

            elif extension == ".pdf":
                if pdf_extract_text:
                    # pdfminer 需要文件路径或文件名，我们将 bytes 写入临时文件
                    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=True) as tmp_file:
                        tmp_file.write(file_bytes)
                        tmp_file_path = tmp_file.name
                        logger.debug(f"将 PDF 内容写入临时文件: {tmp_file_path}")
                        # 确保写入完成
                        tmp_file.flush()
                        # 提取文本
                        content = pdf_extract_text(tmp_file_path) # type: ignore
                        logger.debug("PDF 文本提取完成。")

                else:
                    raise NotSupportedException("未安装 pdfminer.six，无法解析 .pdf 文件。")
            else:
                logger.error(f"不支持的文件类型: {extension}")
                raise NotSupportedException(f"不支持的文件类型: {extension}")

            logger.info(f"文件内容提取完成: Filename='{original_filename}', Extracted Length={len(content)}")
            return content.strip() if content else ""

        except NotSupportedException as e:
            logger.error(f"提取文件内容失败: {e}")
            raise # 直接抛出 NotSupportedException
        except BusinessException as e:
             logger.error(f"提取文件内容时发生业务异常: {e.message}")
             raise # 直接抛出 BusinessException
        except Exception as e:
            logger.error(f"提取文件内容时发生未知错误: URL='{document_url}', Filename='{original_filename}' - {e}")
            raise BusinessException(f"提取文件内容失败: {str(e)}", code=500) from e

    async def extract_web_content_async(self, url: str) -> str:
        """提取网页内容"""
        logger.info(f"开始提取网页内容: URL='{url}'")
        try:
            response = await self._http_client.get(url)
            response.raise_for_status()
            # 读取 bytes 以便后续正确解码
            html_bytes = await response.aread()
            content = self._extract_html_content(html_bytes, response.encoding) # 传递检测到的编码
            logger.info(f"网页内容提取完成: URL='{url}', Extracted Length={len(content)}")
            return content.strip() if content else ""
        except httpx.RequestError as e:
            logger.error(f"提取网页时请求错误: {e.request.url!r} - {e}")
            raise BusinessException(f"无法访问网页 (请求错误): {url}", code=400) from e
        except httpx.HTTPStatusError as e:
            logger.error(f"提取网页时 HTTP 状态错误: {e.request.url!r} - Status {e.response.status_code}")
            raise BusinessException(f"无法访问网页 (HTTP {e.response.status_code}): {url}", code=400) from e
        except Exception as e:
            logger.error(f"提取网页内容时发生未知错误: URL='{url}' - {e}")
            raise BusinessException(f"提取网页内容失败: {str(e)}", code=500) from e

    def _extract_html_content(self, html_bytes: bytes, detected_encoding: Optional[str] = None) -> str:
        """从 HTML 字节中提取主要文本内容"""
        if not BeautifulSoup:
            logger.warning("BeautifulSoup4 未安装，将使用简单的正则去除标签。")
            # 尝试解码
            html_string = ""
            encodings = [detected_encoding] if detected_encoding else []
            encodings.extend(['utf-8', 'gbk', 'gb2312'])
            for enc in encodings:
                 if not enc: continue
                 try:
                      html_string = html_bytes.decode(enc)
                      logger.debug(f"HTML 使用 {enc} 解码成功。")
                      break
                 except UnicodeDecodeError:
                      continue
            if not html_string:
                 html_string = html_bytes.decode('utf-8', errors='ignore')
                 logger.warning("无法确定 HTML 编码，强制 UTF-8 解码。")

            # 简单的正则去除标签
            text = re.sub(r'<script.*?>.*?</script>', '', html_string, flags=re.IGNORECASE | re.DOTALL)
            text = re.sub(r'<style.*?>.*?</style>', '', text, flags=re.IGNORECASE | re.DOTALL)
            text = re.sub(r'<.*?>', ' ', text) # 去除所有标签，替换为空格
            text = re.sub(r'\s+', ' ', text)    # 合并多个空白
            return text.strip()

        # 使用 BeautifulSoup 解析
        try:
            # BeautifulSoup 可以自动检测编码，或者指定 from_encoding
            soup = BeautifulSoup(html_bytes, 'html5lib', from_encoding=detected_encoding)

            # 移除不需要的标签
            for element in soup(["script", "style", "noscript", "iframe", "svg", "nav", "footer", "aside", "form", "button"]):
                element.decompose()

            # 获取 body 或 article 或 main 内容
            main_content = soup.find('article') or soup.find('main') or soup.body
            if main_content:
                text = main_content.get_text(separator=' ', strip=True)
            else:
                 text = soup.get_text(separator=' ', strip=True) # Fallback

            # 进一步清理空白
            text = re.sub(r'\s+', ' ', text).strip()
            return text
        except Exception as e:
             logger.error(f"使用 BeautifulSoup 解析 HTML 时出错: {e}")
             # Fallback 到简单正则
             html_string = html_bytes.decode('utf-8', errors='ignore')
             text = re.sub(r'<script.*?>.*?</script>', '', html_string, flags=re.IGNORECASE | re.DOTALL)
             text = re.sub(r'<style.*?>.*?</style>', '', text, flags=re.IGNORECASE | re.DOTALL)
             text = re.sub(r'<.*?>', ' ', text)
             text = re.sub(r'\s+', ' ', text)
             return text.strip()

    def _extract_docx_content(self, file_stream: io.BytesIO) -> str:
        """从 DOCX 文件流中提取文本"""
        if not docx: return "[错误：未安装 python-docx]"
        try:
            document = docx.Document(file_stream)
            full_text = []
            for para in document.paragraphs:
                full_text.append(para.text)
            # 可以考虑提取表格内容 (更复杂)
            # for table in document.tables:
            #     for row in table.rows:
            #         row_text = [cell.text for cell in row.cells]
            #         full_text.append("\t".join(row_text))
            logger.debug("DOCX 文本提取完成。")
            return '\n'.join(full_text)
        except Exception as e:
            logger.error(f"提取 DOCX 内容失败: {e}")
            raise BusinessException(f"解析 DOCX 文件失败: {str(e)}") from e

    # PDF Miner 提取 (如果使用 pdfminer.six 库)
    # def _extract_pdf_content_pdfminer(self, file_stream: io.BytesIO) -> str:
    #     if not PDFPage: return "[错误：未安装 pdfminer.six]"
    #     try:
    #         resource_manager = PDFResourceManager()
    #         fake_file_handle = StringIO()
    #         # 使用二进制模式读取可能会更好？
    #         # converter = TextConverter(resource_manager, fake_file_handle, laparams=LAParams())
    #         converter = TextConverter(resource_manager, fake_file_handle, codec='utf-8', laparams=LAParams()) # 指定编码
    #         page_interpreter = PDFPageInterpreter(resource_manager, converter)

    #         for page in PDFPage.get_pages(file_stream,
    #                                       caching=True,
    #                                       check_extractable=True):
    #             page_interpreter.process_page(page)

    #         text = fake_file_handle.getvalue()
    #         converter.close()
    #         fake_file_handle.close()
    #         logger.debug("PDF (pdfminer) 文本提取完成。")
    #         return text
    #     except Exception as e:
    #         logger.error(f"使用 pdfminer 提取 PDF 内容失败: {e}")
    #         raise BusinessException(f"解析 PDF 文件失败: {str(e)}") from e